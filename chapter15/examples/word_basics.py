import docx
import read_docx

# Reading Word documents
doc = docx.Document('demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text
doc.paragraphs[1].runs[4].text

# Getting the full text from a docx file (don't forget to import the function)
print(read_docx.get_text('demo.docx'))

# Styling
doc.paragraphs[0].text
doc.paragraphs[0].style # Show current style
doc.paragraphs[0].style = 'Normal' # Change style for the entire paragraph
doc.paragraphs[1].text
doc.paragraphs[1].runs[0].style = 'QuoteChar' # Change style for a run
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx') # Save document to a new docx file

# Writing Word Documents
doc = docx.Document() # Create a new document
doc.add_paragraph('Hello, world!') # Add a new paragraph
doc.save('helloworld.docx')

doc = docx.Document()
doc.add_paragraph('Hello, world!', 'Title') # Second parameter is style
para_obj1 = doc.add_paragraph('This is a second paragraph.')
para_obj2 = doc.add_paragraph('This is a yet another paragraph.')
# add a run to an existing paragraph
para_obj1.add_run(' This text is being added to the second paragraph.', 'Quote Char')
doc.save('multiple_paragraphs.docx')

# Adding Headings
doc = docx.Document()
doc.add_heading('Header 0 text', 0)
doc.add_heading('Header text 1', 1)
doc.add_heading('Header text 2', 2)
doc.add_heading('Header text 3', 3)
doc.add_heading('Header text 4', 4)
doc.save('headings.docx')

# Adding Line and Page braks
doc = docx.Document()
doc.add_paragraph('This is the break document')
doc.paragraphs[0].runs[0].add_break()
doc.paragraphs[0].runs[0].add_break()
doc.add_paragraph('This is on the first page.')
doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('two_page.docx')

# Adding pictures
doc = docx.Document()
doc.add_paragraph('This is the picture document')
# check docx.shared for all kinds of units to pass to width and height
doc.add_picture(
    'zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
doc.save('cat.docx')
