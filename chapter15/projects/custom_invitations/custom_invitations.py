"""
Custom Invitations as Word Documents

The file guests.txt is a list of guest names with one name per line.

Write a program that would generate a Word document with custom invitations that look like:

It would be a pleasure to have the company of (Style1)
GUEST_NAME (Style2)
at 11010 Memory Lane on the Evening of (Style1)
April 1st (Style3)
at 7 o'clock (Style1)
[all text should be centered btw]

Since Python-Docx can use only those styles that already exist in the
Word document, you will have to first add these styles to a blank Word
file and then open that file with Python-Docx.
There should be one invitation per page in the resulting Word document,
so call add_break() to add a page break after the last paragraph of each
invitation. This way, you will need to open only one Word document to
print all of the invitations at once.
"""

import docx

# Load guests into a list
with open('guests.txt') as file_obj:
    guest_list = file_obj.read().split('\n')

# Create word document
doc = docx.Document()

# Create a unique invitation for each guest in the list
for guest in guest_list:
    para = doc.add_paragraph()
    para.alignment = 1
    run = para.add_run('It would be a pleasure to have the company of')
    run.style = 'Quote Char'
    run.add_break()
    run = para.add_run(guest)
    run.style = 'Heading 1 Char'
    run.add_break()
    run = para.add_run('at 11010 Memory Lane on the Evening of')
    run.style = 'Quote Char'
    run.add_break()
    run = para.add_run('April 1st')
    run.style = 'Heading 2 Char'
    run.add_break()
    run = para.add_run("at 7 o'clock")
    run.style = 'Quote Char'
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

# Save the custom invitations
doc.save('custom_invitations.docx')    
